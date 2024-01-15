import os
import random
from docx.shared import Cm, Pt
from docx import Document

spiskiDirectory = '/workspaces/Projekts/Spiski'
protokoliDirectory = '/workspaces/Projekts/Protokoli'

def createWordFromText(textFilePath):
    with open(textFilePath, 'r') as file:
        content = file.readlines()

    random.shuffle(content)

    doc = Document()

    headingTxt = os.path.splitext(os.path.basename(textFilePath))[0]
    heading = doc.add_heading(level=1)
    run = heading.add_run(headingTxt)
    run.font.size = Pt(30)

    table = doc.add_table(rows=0, cols=4, style='TableGrid')

    for index, line in enumerate(content):
        values = line.strip().split(',')

        row = table.add_row().cells

        row[1].text = str(index + 1)  
        row[2].text = values[0].strip()  
        row[3].text = values[1].strip() 
     

    for i in range(0, 4):
        for col in table.columns[i].cells:
            col.width = Cm(0.8)
            for paragraph in col.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(24)
          
   

    filename = os.path.splitext(os.path.basename(textFilePath))[0] + '.docx'
    doc.save(os.path.join(protokoliDirectory, filename))

textFiles = os.listdir(spiskiDirectory)

for filename in textFiles:
    textFilePath = os.path.join(spiskiDirectory, filename)
    createWordFromText(textFilePath)
